"""Esportazione del testo LETTERALE in .txt e .docx.

Il livello letterale e' l'artefatto primario. La colonna "Interlocutore" (speaker) usa le
etichette assegnate (diarizzazione + correzione manuale); resta vuota se non assegnate.
I timestamp sono opzionali.
"""
from __future__ import annotations

import io


def _fmt_ts(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    return f"{m:02d}:{s:02d}"


def to_txt(project: dict, with_timestamps: bool = False) -> str:
    lines: list[str] = [f"Trascrizione letterale — {project.get('name', '')}", ""]
    for seg in project["segments"]:
        prefix = f"[{_fmt_ts(seg['start'])}] " if with_timestamps else ""
        speaker = seg.get("speaker", "")
        label = f"{speaker}:\t" if speaker else "\t"  # tab = colonna interlocutore
        lines.append(f"{prefix}{label}{seg['literal']}")
    return "\n".join(lines) + "\n"


def to_docx_bytes(project: dict, with_timestamps: bool = False) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(f"Trascrizione letterale — {project.get('name', '')}", level=1)

    cols = 3 if with_timestamps else 2
    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    idx = 0
    if with_timestamps:
        hdr[idx].text = "Tempo"; idx += 1
    hdr[idx].text = "Interlocutore"; idx += 1
    hdr[idx].text = "Testo (letterale)"

    for seg in project["segments"]:
        row = table.add_row().cells
        idx = 0
        if with_timestamps:
            row[idx].text = _fmt_ts(seg["start"]); idx += 1
        row[idx].text = seg.get("speaker", "")   # etichetta interlocutore (vuota se assente)
        idx += 1
        run = row[idx].paragraphs[0].add_run(seg["literal"])
        run.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
